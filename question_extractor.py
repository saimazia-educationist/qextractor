from docx.shared import Pt
import pdfplumber
import re
import logging
import pytesseract
import shutil
import fitz  # PyMuPDF
import os, io
from collections import defaultdict
from docx.shared import Inches
from docx import Document


class InvalidPaperError(Exception):
    """Raised when a PDF's filename doesn't match a recognized paper code."""
    pass


def _configure_tesseract():
    """
    Find the Tesseract binary without hardcoding a Windows path.
    Priority: TESSERACT_CMD env var -> already on PATH -> common install locations.
    """
    env_path = os.environ.get("TESSERACT_CMD")
    if env_path and os.path.exists(env_path):
        pytesseract.pytesseract.tesseract_cmd = env_path
        return

    found = shutil.which("tesseract")
    if found:
        pytesseract.pytesseract.tesseract_cmd = found
        return

    for candidate in (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
        "/opt/homebrew/bin/tesseract",
    ):
        if os.path.exists(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate
            return

    logging.warning(
        "Tesseract binary not found. Set the TESSERACT_CMD environment "
        "variable to its full path, or install it and ensure it's on PATH."
    )


_configure_tesseract()

global new_question, add_question_num, add_answer_no, paper_label
logging.basicConfig(level=logging.INFO)
paper_map = {
    "11": "Paper 1",
    "12": "Paper 1",
    "13": "Paper 1",
    "21": "Paper 2",
    "22": "Paper 2",
    "23": "Paper 2",
    "31": "Paper 3",
    "32": "Paper 3",
    "33": "Paper 3",
    "41": "Paper 4",
    "42": "Paper 4",
    "43": "Paper 4",

}
keyword_map1 = {
    "editing software": "Information representation.pdf",
    "feature": "Information representation.pdf",
    "analogue": "Information representation.pdf",
    "sample": "Information representation.pdf",
    "digital": "Information representation.pdf",
    "sampling resolution": "Information representation.pdf",
    "number system": "Information representation.pdf",
    "binary addition": "Information representation.pdf",
    "overflow": "Information representation.pdf",
    "mebibytes": "Information representation.pdf",
    "hexadecimal value ": "Information representation.pdf",
    "Binary Coded Decimal": "Information representation.pdf",
    "two's complement": "Information representation.pdf",
    "denary": "Information representation.pdf",
    "Unicode": "Information representation.pdf",
    "Character set": "Information representation.pdf",
    "sampling": "Information representation.pdf",
    "resolution": "Information representation.pdf",
    "editing": "Information representation.pdf",
    "sound": "Information representation.pdf",
    "compression": "Information representation.pdf",
    "lossy": "Information representation.pdf",
    "lossless": "Information representation.pdf",
    "run-length encoding": "Information representation.pdf",
    "rle": "Information representation.pdf",
    "sampling rate": "Information representation.pdf",
    "file size": "Information representation.pdf",
    "bitmap": "Information representation.pdf",
    "binary": "Information representation.pdf",
    "hexadecimal": "Information representation.pdf",
    "two’s complement": "Information representation.pdf",
    "bcd": "Information representation.pdf",
    "image": "Information representation.pdf",
    "compress": "Information representation.pdf",
    "pixel": "Information representation.pdf",
    "kilo": "Information representation.pdf",
    "vector": "Information representation.pdf",
    "graphic": "Information representation.pdf",
    "gibi": "Information representation.pdf",

    "url": "Communication and networking technologies.pdf",
    "dns": "Communication and networking technologies.pdf",
    "isp": "Communication and networking technologies.pdf",
    "cookies": "Communication and networking technologies.pdf",
    "tcp/io": "Communication and networking technologies.pdf",
    "serial transmission": "Communication and networking technologies.pdf",
    "twisted pair": "Communication and networking technologies.pdf",
    "coaxial": "Communication and networking technologies.pdf",
    "fiber optic": "Communication and networking technologies.pdf",
    "wireless transmission": "Communication and networking technologies.pdf",
    "symmetric encryption": "Communication and networking technologies.pdf",
    "firewall": "Communication and networking technologies.pdf",
    "proxy server": "Communication and networking technologies.pdf",
    "https": "Communication and networking technologies.pdf",
    "ssl/tls": "Communication and networking technologies.pdf",
    "handshaking": "Communication and networking technologies.pdf",
    "lan": "Communication and networking technologies.pdf",
    "wan": "Communication and networking technologies.pdf",
    "io address": "Communication and networking technologies.pdf",
    "mac address": "Communication and networking technologies.pdf",
    "protocol": "Communication and networking technologies.pdf",
    "data packet": "Communication and networking technologies.pdf",
    "router": "Communication and networking technologies.pdf",
    "switch": "Communication and networking technologies.pdf",
    "topology": "Communication and networking technologies.pdf",
    "pstn": "Communication and networking technologies.pdf",
    "star": "Communication and networking technologies.pdf",
    "packets": "Communication and networking technologies.pdf",
    "serial": "Communication and networking technologies.pdf",
    "simplex": "Communication and networking technologies.pdf",
    "duplex": "Communication and networking technologies.pdf",
    "private": "Communication and networking technologies.pdf",
    "ip": "Communication and networking technologies.pdf",
    "world wide web": "Communication and networking technologies.pdf",

    "fuser": "Hardware.pdf",
    "toner": "Hardware.pdf",
    "drum": "Hardware.pdf",
    "secondary storage": "Hardware.pdf",
    "electrostatic": "Hardware.pdf",
    "electrical charge": "Hardware.pdf",
    "storage": "Hardware.pdf",
    "dram": "Hardware.pdf",
    "sram": "Hardware.pdf",
    "storage devices": "Hardware.pdf",
    "optical media": "Hardware.pdf",
    "dynamic": "Hardware.pdf",
    "static": "Hardware.pdf",
    "cache": "Hardware.pdf",
    "dvd": "Hardware.pdf",
    "flash": "Hardware.pdf",
    "non-volatile memory": "Hardware.pdf",
    "lcd": "Hardware.pdf",
    "led": "Hardware.pdf",
    "touch screen": "Hardware.pdf",
    "resistive technology": "Hardware.pdf",
    "capacitive technology": "Hardware.pdf",
    "infrared technology": "Hardware.pdf",
    "barcode reader": "Hardware.pdf",
    "scanner": "Hardware.pdf",
    "microphone": "Hardware.pdf",
    "interactive whiteboard": "Hardware.pdf",
    "digital light projector": "Hardware.pdf",
    "inkjet": "Hardware.pdf",
    "stepper motor": "Hardware.pdf",
    "print head": "Hardware.pdf",
    "disk": "Hardware.pdf",
    "drive": "Hardware.pdf",
    "laser": "Hardware.pdf",
    "external storage": "Hardware.pdf",
    "printer": "Hardware.pdf",
    "3d printer": "Hardware.pdf",
    "ssd": "Hardware.pdf",
    "hdd": "Hardware.pdf",
    "ram": "Hardware.pdf",
    "rom": "Hardware.pdf",
    "off-line": "Hardware.pdf",
    "primary": "Hardware.pdf",
    "secondary": "Hardware.pdf",
    "interrupt signal": "Hardware.pdf",
    "concept keyboard": "Hardware.pdf",
    "arq": "Hardware.pdf",
    "blue laser": "Hardware.pdf",
    "red laser": "Hardware.pdf",
    "2d scanner": "Hardware.pdf",
    "3d scanner": "Hardware.pdf",
    "led display": "Hardware.pdf",
    "lcd projector": "Hardware.pdf",
    "2d cutter": "Hardware.pdf",

    "logic gates": "Logic gates and logic circuits.pdf",
    "and gate": "Logic gates and logic circuits.pdf",
    "or gate": "Logic gates and logic circuits.pdf",
    "not gate": "Logic gates and logic circuits.pdf",
    "nand gate": "Logic gates and logic circuits.pdf",
    "nor gate": "Logic gates and logic circuits.pdf",
    "xor gate": "Logic gates and logic circuits.pdf",
    "truth table": "Logic gates and logic circuits.pdf",
    "logic circuit": "Logic gates and logic circuits.pdf",
    "logic statement": "Logic gates and logic circuits.pdf",
    "alarm system": "Logic gates and logic circuits.pdf",
    "x": "Logic gates and logic circuits.pdf",
    "working space": "Logic gates and logic circuits.pdf",
    "description of condition": "Logic gates and logic circuits.pdf",
    "condition": "Logic gates and logic circuits.pdf",

    "von neumann": " Processor fundamentals.pdf",
    "central processing unit": " Processor fundamentals.pdf",
    "mar": " Processor fundamentals.pdf",
    "mdr": " Processor fundamentals.pdf",
    "cir": " Processor fundamentals.pdf",
    "arithmetic and logic": " Processor fundamentals.pdf",
    "alu": " Processor fundamentals.pdf",
    "registers": " Processor fundamentals.pdf",
    "accumulator": " Processor fundamentals.pdf",
    "system bus": " Processor fundamentals.pdf",
    "data bus": " Processor fundamentals.pdf",
    "address bus": " Processor fundamentals.pdf",
    "control bus": " Processor fundamentals.pdf",
    "bus width": " Processor fundamentals.pdf",
    "i/o port": " Processor fundamentals.pdf",
    "usb": " Processor fundamentals.pdf",
    "fetch": " Processor fundamentals.pdf",
    "decode execute cycle": " Processor fundamentals.pdf",
    "register transfer notation": " Processor fundamentals.pdf",
    "interrupt handling": " Processor fundamentals.pdf",
    "interrupt service routine": " Processor fundamentals.pdf",
    "ide": " Processor fundamentals.pdf",
    "features": " Processor fundamentals.pdf",
    "debugging": " Processor fundamentals.pdf",
    "high level language": " Processor fundamentals.pdf",
    "low level language": " Processor fundamentals.pdf",
    "library program": " Processor fundamentals.pdf",
    "clock speed": " Processor fundamentals.pdf",
    "cores": " Processor fundamentals.pdf",
    "mathematical operations": " Processor fundamentals.pdf",
    "comparison": " Processor fundamentals.pdf",
    "performance": " Processor fundamentals.pdf",
    "electronic pulse": " Processor fundamentals.pdf",
    "logical": " Processor fundamentals.pdf",
    "carries data": " Processor fundamentals.pdf",
    "signal": " Processor fundamentals.pdf",

    "translate": "Assembly language programming.pdf",
    "two pass": "Assembly language programming.pdf",
    "pass": "Assembly language programming.pdf",
    "assembly language": "Assembly language programming.pdf",
    "comments": "Assembly language programming.pdf",
    "mode": "Assembly language programming.pdf",
    "out": "Assembly language programming.pdf",
    "instruction": "Assembly language programming.pdf",
    "trace table": "Assembly language programming.pdf",
    "symbol table": "Assembly language programming.pdf",
    "opcode": "Assembly language programming.pdf",
    "operand": "Assembly language programming.pdf",
    "machine code": "Assembly language programming.pdf",
    "assembler": "Assembly language programming.pdf",
    "macros": "Assembly language programming.pdf",
    "directives": "Assembly language programming.pdf",
    "labels": "Assembly language programming.pdf",
    "symbolic": "Assembly language programming.pdf",
    "relative": "Assembly language programming.pdf",
    "absolute": "Assembly language programming.pdf",
    "addressing": "Assembly language programming.pdf",
    "ldd": "Assembly language programming.pdf",
    "cmp": "Assembly language programming.pdf",
    "jpn": "Assembly language programming.pdf",
    "sto": "Assembly language programming.pdf",
    "end": "Assembly language programming.pdf",
    "ldm": "Assembly language programming.pdf",
    "ldi": "Assembly language programming.pdf",
    "add": "Assembly language programming.pdf",
    "sub": "Assembly language programming.pdf",
    "modes": "Assembly language programming.pdf",
    "ldx": "Assembly language programming.pdf",
    "mov": "Assembly language programming.pdf",
    "jmp": "Assembly language programming.pdf",
    "inc": "Assembly language programming.pdf",
    "dec": "Assembly language programming.pdf",
    "logical shift": "Assembly language programming.pdf",
    "lsl": "Assembly language programming.pdf",
    "lsr": "Assembly language programming.pdf",
    "cyclic shift": "Assembly language programming.pdf",
    "arithmetic shift": "Assembly language programming.pdf",
    "<address>": "Assembly language programming.pdf",
    "<register>": "Assembly language programming.pdf",
    "memory address": "Assembly language programming.pdf",
    "instruction address": "Assembly language programming.pdf",

    "monitoring": "Monitoring and control systems.pdf",
    "control system": "Monitoring and control systems.pdf",
    "sensor": "Monitoring and control systems.pdf",
    "actuator": "Monitoring and control systems.pdf",
    "feedback": "Monitoring and control systems.pdf",
    "bit manipulation": "Monitoring and control systems.pdf",
    "bit": "Monitoring and control systems.pdf",
    "bitwise operation": "Monitoring and control systems.pdf",
    "toggle": "Monitoring and control systems.pdf",
    "flag": "Monitoring and control systems.pdf",
    "bit position": "Monitoring and control systems.pdf",
    "green house": "Monitoring and control systems.pdf",
    "intruder": "Monitoring and control systems.pdf",
    "adc": "Monitoring and control systems.pdf",
    "dac": "Monitoring and control systems.pdf",

    "os": "System software.pdf",
    "printer management": "System software.pdf",
    "utility software": "System software.pdf",
    "peripheral management": "System software.pdf",
    "operating system": "System software.pdf",
    "graphical user interface": "System software.pdf",
    "command-line interface": "System software.pdf",
    "resource management": "System software.pdf",
    "memory management": "System software.pdf",
    "scheduling": "System software.pdf",
    "device management": "System software.pdf",
    "file management": "System software.pdf",
    "security management": "System software.pdf",
    "utility": "System software.pdf",
    "disk formatter": "System software.pdf",
    "checker": "System software.pdf",
    "defragmenter": "System software.pdf",
    "backup software": "System software.pdf",
    "file compression": "System software.pdf",
    "virus checker": "System software.pdf",
    "program libraries": "System software.pdf",
    "dynamic link library": "System software.pdf",
    "dll": "System software.pdf",
    "translators": "System software.pdf",
    "compiler": "System software.pdf",
    "interpreter": "System software.pdf",
    "interpreted": "System software.pdf",
    "compiled": "System software.pdf",
    "prettyprinting": "System software.pdf",
    "context-sensitive prompts": "System software.pdf",
    "dynamic syntax checks": "System software.pdf",
    "integrated development environment": "System software.pdf",

    "data block": "Security, privacy and data integrity.pdf",
    "calculate": "Security, privacy and data integrity.pdf",
    "data integrity.": "Security, privacy and data integrity.pdf",
    "data privacy": "Security, privacy and data integrity.pdf",
    "security": "Security, privacy and data integrity.pdf",
    "threats": "Security, privacy and data integrity.pdf",
    "malware": "Security, privacy and data integrity.pdf",
    "virus": "Security, privacy and data integrity.pdf",
    "worm": "Security, privacy and data integrity.pdf",
    "bot": "Security, privacy and data integrity.pdf",
    "logic bomb": "Security, privacy and data integrity.pdf",
    "spyware": "Security, privacy and data integrity.pdf",
    "phishing": "Security, privacy and data integrity.pdf",
    "pharming": "Security, privacy and data integrity.pdf",
    "keylogger": "Security, privacy and data integrity.pdf",
    "password": "Security, privacy and data integrity.pdf",
    "update": "Security, privacy and data integrity.pdf",
    "anti virus": "Security, privacy and data integrity.pdf",
    "intrusion": "Security, privacy and data integrity.pdf",
    "data loss": "Security, privacy and data integrity.pdf",
    "backup": "Security, privacy and data integrity.pdf",
    "restrict": "Security, privacy and data integrity.pdf",
    "access": "Security, privacy and data integrity.pdf",
    "authorisation": "Security, privacy and data integrity.pdf",
    "authentication": "Security, privacy and data integrity.pdf",
    "validation": "Security, privacy and data integrity.pdf",
    "verification": "Security, privacy and data integrity.pdf",
    "check digit": "Security, privacy and data integrity.pdf",
    "data transfer": "Security, privacy and data integrity.pdf",
    "parity bit": "Security, privacy and data integrity.pdf",
    "parity byte": "Security, privacy and data integrity.pdf",
    "unauthorised": "Security, privacy and data integrity.pdf",
    "policy": "Security, privacy and data integrity.pdf",
    "measures": "Security, privacy and data integrity.pdf",
    "checksum": "Security, privacy and data integrity.pdf",
    "authentic": "Security, privacy and data integrity.pdf",
    "even parity": "Security, privacy and data integrity.pdf",
    "odd parity": "Security, privacy and data integrity.pdf",
    "format check": "Security, privacy and data integrity.pdf",
    "range check": "Security, privacy and data integrity.pdf",
    "double entry": "Security, privacy and data integrity.pdf",
    "digital signature": "Security, privacy and data integrity.pdf",
    "protect": "Security, privacy and data integrity.pdf",
    "back-up": "Security, privacy and data integrity.pdf",

    "Ethics": "Ethics and ownership.pdf",
    "profession": "Ethics and ownership.pdf",
    "Public": "Ethics and ownership.pdf",
    "client": "Ethics and ownership.pdf",
    "employer": "Ethics and ownership.pdf",
    "product": "Ethics and ownership.pdf",
    "judgement": "Ethics and ownership.pdf",
    "management": "Ethics and ownership.pdf",
    "colleagues": "Ethics and ownership.pdf",
    "self": "Ethics and ownership.pdf",
    "public good": "Ethics and ownership.pdf",
    "ownership": "Ethics and ownership.pdf",
    "copyright": "Ethics and ownership.pdf",
    "license": "Ethics and ownership.pdf",
    "commercial": "Ethics and ownership.pdf",
    "freeware": "Ethics and ownership.pdf",
    "shareware": "Ethics and ownership.pdf",
    "open source": "Ethics and ownership.pdf",
    "free software": "Ethics and ownership.pdf",
    "artificial intelligence": "Ethics and ownership.pdf",
    "ai": "Ethics and ownership.pdf",
    "problem solving": "Ethics and ownership.pdf",
    "perception": "Ethics and ownership.pdf",
    "reasoning": "Ethics and ownership.pdf",
    "impact of ai": "Ethics and ownership.pdf",
    "robots": "Ethics and ownership.pdf",
    "ethical": "Ethics and ownership.pdf",
    "unethical": "Ethics and ownership.pdf",
    "acm/ieee": "Ethics and ownership.pdf",
    "software engineer": "Ethics and ownership.pdf",
    "company": "Ethics and ownership.pdf",
    "principle": "Ethics and ownership.pdf",
    "licence": "Ethics and ownership.pdf",
    "trial": "Ethics and ownership.pdf",
    "distributed": "Ethics and ownership.pdf",
    "software company": "Ethics and ownership.pdf",

    "relational": "Databases.pdf",
    "table": "Databases.pdf",
    "normal form": "Databases.pdf",
    "normalize": "Databases.pdf",
    "tuple": "Databases.pdf",
    "sql": "Databases.pdf",
    "data redundancy": "Databases.pdf",
    "attribute": "Databases.pdf",
    "candidate key": "Databases.pdf",
    "secondary key:": "Databases.pdf",
    "foreign key": "Databases.pdf",
    "referential integrity": "Databases.pdf",
    "(e–r) diagram": "Databases.pdf",
    "entity-relationship": "Databases.pdf",
    "entity": "Databases.pdf",
    "1:1": "Databases.pdf",
    "1:m": "Databases.pdf",
    "m:1": "Databases.pdf",
    "conceptual level": "Databases.pdf",
    "database management system": "Databases.pdf",
    "dbms": "Databases.pdf",
    "database administrator": "Databases.pdf",
    "conceptual schema": "Databases.pdf",
    "query": "Databases.pdf",
    "integrity": "Databases.pdf",
    "data definition language": "Databases.pdf",
    "data manipulation language": "Databases.pdf",
    "select": "Databases.pdf",
    "database": "Databases.pdf",
    "group by": "Databases.pdf",
    "order by": "Databases.pdf",
    "er diagram": "Databases.pdf",
    "record": "Databases.pdf", }

keyword_map2 = {
    "structure chart": "Structure chart.pdf",
    "structure": "Structure chart.pdf",
    "chart": "Structure chart.pdf",
    "modules": "Structure chart.pdf",
    "curved arrow": "Structure chart.pdf",
    "abstraction": "Structure chart.pdf",
    "diamond symbol": "Structure chart.pdf",

    "analysis": "Software development.pdf",
    "design": "Software development.pdf",
    "coding": "Software development.pdf",
    "testing": "Software development.pdf",
    "waterfall": "Software development.pdf",
    "iterative": "Software development.pdf",
    "rapid application": "Software development.pdf",
    "type of error": "Software development.pdf",
    "syntax error": "Software development.pdf",
    "logic": "Software development.pdf",
    "runtime": "Software development.pdf",
    "stub": "Software development.pdf",
    "black box": "Software development.pdf",
    "white box": "Software development.pdf",
    "integration testing": "Software development.pdf",
    "alpha": "Software development.pdf",
    "beta": "Software development.pdf",
    "test data": "Software development.pdf",
    "normal": "Software development.pdf",
    "abnormal": "Software development.pdf",
    "boundary": "Software development.pdf",
    "extreme": "Software development.pdf",
    "maintenance": "Software development.pdf",
    "corrective": "Software development.pdf",
    "adaptive": "Software development.pdf",
    "perfective": "Software development.pdf",
    "integrated development environment": "Software development.pdf",
    "ide": "Software development.pdf",
    "expected outcome": "Software development.pdf",
    "test plan": "Software development.pdf",
    "validate": "Software development.pdf",
    "programming errors": "Software development.pdf",

    "queue": "Abstract data types.pdf",
    "linked list": "Abstract data types.pdf",
    "stack": "Abstract data types.pdf",
    "pointer": "Abstract data types.pdf",

    "record": "Record data type.pdf",

    "current state": "State transition diagrams.pdf",
    "final state": "State transition diagrams.pdf",
    "state": "State transition diagrams.pdf",
    "start state": "State transition diagrams.pdf",
    "transition": "State transition diagrams.pdf",
    "input state": "State transition diagrams.pdf",
    "output state": "State transition diagrams.pdf",
    "next state": "State transition diagrams.pdf",
    "decomposition": "State transition diagrams.pdf",

    "procedure": "Modules & Arrays.pdf",
    "function": "Modules & Arrays.pdf",
    "algorithm": "Modules & Arrays.pdf",
    "returns": "Modules & Arrays.pdf",
    "endfunction": "Modules & Arrays.pdf",
    "structured english": "Modules & Arrays.pdf",
    "endprocedure": "Modules & Arrays.pdf",
    "module": "Modules & Arrays.pdf",
    "array": "Modules & Arrays.pdf",
    "1D": "Modules & Arrays.pdf",
    "2D": "Modules & Arrays.pdf",
    "column": "Modules & Arrays.pdf",
    "row": "Modules & Arrays.pdf",

    "flow chart": "Flow chart.pdf",
    "yes": "Flow chart.pdf",
    "no": "Flow chart.pdf",
    "decision": "Flow chart.pdf",  # diamond shape
    "process": "Flow chart.pdf",  # rectangle
    "terminator": "Flow chart.pdf",  # start/end oval
    "arrow": "Flow chart.pdf",  # connectors

    "text": "Text files.pdf",
    "file": "Text files.pdf",
    ".txt": "Text files.pdf",

    "selection": "Programming constructs and string functions.pdf",
    "iteration": "Programming constructs and string functions.pdf",
    "expression": "Programming constructs and string functions.pdf",
    "myint": "Programming constructs and string functions.pdf",
    "mychar": "Programming constructs and string functions.pdf",
    "char": "Programming constructs and string functions.pdf",
    "string function": "Programming constructs and string functions.pdf",
    "valid": "Programming constructs and string functions.pdf",
    "repetition": "Programming constructs and string functions.pdf",
    "error": "Programming constructs and string functions.pdf",
    "evaluates to": "Programming constructs and string functions.pdf",
    "data types": "Programming constructs and string functions.pdf",
    "evaluate": "Programming constructs and string functions.pdf",
    "case of": "Programming constructs and string functions.pdf",
    "endcase": "Programming constructs and string functions.pdf",
    "complete the table": "Programming constructs and string functions.pdf",
    "answer": "Programming constructs and string functions.pdf",
    "assignment": "Programming constructs and string functions.pdf",
    "min": "Programming constructs and string functions.pdf",
    "max": "Programming constructs and string functions.pdf",
    "str_to_num": "Programming constructs and string functions.pdf",

}

keyword_map3 = {
    "data type": "Data representation.pdf",
    "built-in": "Data representation.pdf",
    "user-defined": "Data representation.pdf",
    "composite": "Data representation.pdf",
    "enumerated": "Data representation.pdf",
    "declare": "Data representation.pdf",
    "ordinal": "Data representation.pdf",
    "record": "Data representation.pdf",
    "type": "Data representation.pdf",
    "class": "Data representation.pdf",
    "pointer ": "Data representation.pdf",
    "caret": "Data representation.pdf",
    "representation": "Data representation.pdf",
    "set": "Data representation.pdf",
    "file organisation": "Data representation.pdf",
    "serial": "Data representation.pdf",
    "approximation": "Data representation.pdf",
    "sequential": "Data representation.pdf",
    "random file": "Data representation.pdf",
    "direct access": "Data representation.pdf",
    "non-composite": "Data representation.pdf",
    "hashing algorithm": "Data representation.pdf",
    "file access": "Data representation.pdf",
    "real number": "Data representation.pdf",
    "floating-point": "Data representation.pdf",
    "fixed point": "Data representation.pdf",
    "largest positive value": "Data representation.pdf",
    "smallest positive value": "Data representation.pdf",
    "largest negative value": "Data representation.pdf",
    "smallest negative value": "Data representation.pdf",
    "mantissa": "Data representation.pdf",
    "exponent": "Data representation.pdf",
    "precision": "Data representation.pdf",
    "two’s complement": "Data representation.pdf",
    "normalised": "Data representation.pdf",
    "definition": "Data representation.pdf",
    "programmer": "Data representation.pdf",
    "statement that assigns": "Data representation.pdf",
    "type definition": "Data representation.pdf",
    "state the current values of the following expressions": "Data representation.pdf",
    "smallest positive number": "Data representation.pdf",
    "largest positive number": "Data representation.pdf",
    "smallest negative number": "Data representation.pdf",
    "largest negative number": "Data representation.pdf",
    "denary value": "Data representation.pdf",
    "range": "Data representation.pdf",

    "packet": "Communication and Internet technologies.pdf",
    "circuit": "Communication and Internet technologies.pdf",
    "switching": "Communication and Internet technologies.pdf",
    "router": "Communication and Internet technologies.pdf",
    "communication": "Communication and Internet technologies.pdf",
    "video conferencing": "Communication and Internet technologies.pdf",
    "tcp/ip": "Communication and Internet technologies.pdf",
    "stack": "Communication and Internet technologies.pdf",
    "layer": "Communication and Internet technologies.pdf",
    "transport layer": "Communication and Internet technologies.pdf",
    "internet layer": "Communication and Internet technologies.pdf",
    "data link": "Communication and Internet technologies.pdf",
    "protocol": "Communication and Internet technologies.pdf",
    "application layer": "Communication and Internet technologies.pdf",
    "bittorrent": "Communication and Internet technologies.pdf",
    "smtp": "Communication and Internet technologies.pdf",
    "imap": "Communication and Internet technologies.pdf",
    "p2p": "Communication and Internet technologies.pdf",
    "peer-to-peer": "Communication and Internet technologies.pdf",
    "routing table": "Communication and Internet technologies.pdf",
    "data packet": "Communication and Internet technologies.pdf",
    "csma/cd": "Communication and Internet technologies.pdf",
    "tcp": "Communication and Internet technologies.pdf",
    "ip": "Communication and Internet technologies.pdf",
    "star": "Communication and Internet technologies.pdf",
    "topology": "Communication and Internet technologies.pdf",
    "bus": "Communication and Internet technologies.pdf",
    "ring": "Communication and Internet technologies.pdf",
    "mesh": "Communication and Internet technologies.pdf",
    "http": "Communication and Internet technologies.pdf",

    "simd": "Hardware and virtual machines.pdf",
    "mimd": "Hardware and virtual machines.pdf",
    "misd": "Hardware and virtual machines.pdf",
    "sisd": "Hardware and virtual machines.pdf",
    "computer architecture": "Hardware and virtual machines.pdf",
    "massively parallel": "Hardware and virtual machines.pdf",
    "cisc": "Hardware and virtual machines.pdf",
    "risc": "Hardware and virtual machines.pdf",
    "pipelining": "Hardware and virtual machines.pdf",
    "fetch": "Hardware and virtual machines.pdf",
    "decode": "Hardware and virtual machines.pdf",
    "execute": "Hardware and virtual machines.pdf",
    "fixed-length": "Hardware and virtual machines.pdf",
    "variable-length": "Hardware and virtual machines.pdf",
    "instruction-level": "Hardware and virtual machines.pdf",
    "parallelism": "Hardware and virtual machines.pdf",
    "interrupt handling": "Hardware and virtual machines.pdf",
    "virtual machine": "Hardware and virtual machines.pdf",
    "host": "Hardware and virtual machines.pdf",
    "guest": "Hardware and virtual machines.pdf",
    "operating system": "Hardware and virtual machines.pdf",
    "write back": "Hardware and virtual machines.pdf",

    "half adder": "Logic circuits and Boolean algebra.pdf",
    "full adder": "Logic circuits and Boolean algebra.pdf",
    "sequential logic circuit": "Logic circuits and Boolean algebra.pdf",
    "combinational circuit": "Logic circuits and Boolean algebra.pdf",
    "flip-flop": "Logic circuits and Boolean algebra.pdf",
    "jk": "Logic circuits and Boolean algebra.pdf",
    "boolean algebra": "Logic circuits and Boolean algebra.pdf",
    "de morgan": "Logic circuits and Boolean algebra.pdf",
    "k-map": "Logic circuits and Boolean algebra.pdf",
    "truth table": "Logic circuits and Boolean algebra.pdf",
    "sum-of-products": "Logic circuits and Boolean algebra.pdf",
    "draw loop(s)": "Logic circuits and Boolean algebra.pdf",
    "groups": "Logic circuits and Boolean algebra.pdf",
    "boolean function": "Logic circuits and Boolean algebra.pdf",
    "nor gate": "Logic circuits and Boolean algebra.pdf",
    "nand gate": "Logic circuits and Boolean algebra.pdf",
    "logic circuit": "Logic circuits and Boolean algebra.pdf",
    "karnaugh": "Logic circuits and Boolean algebra.pdf",
    "state": "Logic circuits and Boolean algebra.pdf",
    "sr": "Logic circuits and Boolean algebra.pdf",
    "boolean expression": "Logic circuits and Boolean algebra.pdf",

    "process": "System Software Advanced.pdf",
    "running": "System Software Advanced.pdf",
    "ready": "System Software Advanced.pdf",
    "blocked": "System Software Advanced.pdf",
    "high-level scheduler": "System Software Advanced.pdf",
    "low-level scheduler": "System Software Advanced.pdf",
    "shortest job first": "System Software Advanced.pdf",
    "round robin": "System Software Advanced.pdf",
    "first come first served": "System Software Advanced.pdf",
    "scheduling": "System Software Advanced.pdf",
    "process management": "System Software Advanced.pdf",
    "rpn": "System Software Advanced.pdf",
    "reverse polish notation": "System Software Advanced.pdf",
    "": "System Software Advanced.pdf",
    "infix": "System Software Advanced.pdf",
    "prefix": "System Software Advanced.pdf",
    "scheduling algorithm": "System Software Advanced.pdf",
    "memory management": "System Software Advanced.pdf",
    "segment": "System Software Advanced.pdf",
    "paging": "System Software Advanced.pdf",
    "virtual memory": "System Software Advanced.pdf",
    "disk thrashing": "System Software Advanced.pdf",
    "translation software": "System Software Advanced.pdf",
    "lexical": "System Software Advanced.pdf",
    "syntax analysis": "System Software Advanced.pdf",
    "semantic": "System Software Advanced.pdf",
    "intermediate code": "System Software Advanced.pdf",
    "symbol table": "System Software Advanced.pdf",
    "parse tree": "System Software Advanced.pdf",
    "syntax tree": "System Software Advanced.pdf",
    "syntax diagram": "System Software Advanced.pdf",
    "grammar": "System Software Advanced.pdf",
    "back-end synthesis": "System Software Advanced.pdf",

    "encrypt": "Security.pdf",
    "key cryptography": "Security.pdf",
    "quantum": "Security.pdf",
    "cryptography": "Security.pdf",
    "digital signature": "Security.pdf",
    "certificate": "Security.pdf",
    "private key": "Security.pdf",
    "public key": "Security.pdf",
    "symmetric": "Security.pdf",
    "asymmetric": "Security.pdf",
    "secure socket layer": "Security.pdf",
    "ssl": "Security.pdf",
    "tls": "Security.pdf",
    "transport layer security": "Security.pdf",
    "protocol": "Security.pdf",
    "decrypt": "Security.pdf",

    "graphs": "artificial intelligence.pdf",
    "ai": "artificial intelligence.pdf",
    "dijkistra's": "artificial intelligence.pdf",
    "shortest path": "artificial intelligence.pdf",
    "a*": "artificial intelligence.pdf",
    "node": "artificial intelligence.pdf",
    "machine learning": "artificial intelligence.pdf",
    "supervised": "artificial intelligence.pdf",
    "unsupervised": "artificial intelligence.pdf",
    "reinforcement": "artificial intelligence.pdf",
    "regression analysis": "artificial intelligence.pdf",
    "neural network": "artificial intelligence.pdf",
    "deep learning": "artificial intelligence.pdf",
    "shortest time": "artificial intelligence.pdf",
    "heuristic": "artificial intelligence.pdf",
    "hidden layer": "artificial intelligence.pdf",

    "recursively": "Recursion.pdf",
    "recursion": "Recursion.pdf",
    "recursive": "Recursion.pdf",
    "iterative": "Recursion.pdf",
    "base case": "Recursion.pdf",
    "general case": "Recursion.pdf",

    "object": "OOP.pdf",
    "class": "OOP.pdf",
    "private": "OOP.pdf",
    "public": "OOP.pdf",
    "inheritance": "OOP.pdf",
    "aggregation": "OOP.pdf",
    "containment": "OOP.pdf",
    "OOP.pdf": "OOP.pdf",
    "attribute": "OOP.pdf",
    "method": "OOP.pdf",
    "encapsulation": "OOP.pdf",
    "getter": "OOP.pdf",
    "setter": "OOP.pdf",
    "constructor": "OOP.pdf",
    "instantiate": "OOP.pdf",
    "abstract": "OOP.pdf",
    "derived class": "OOP.pdf",
    "polymorphism": "OOP.pdf",

    "instruction set": "Low level programming.pdf",
    "op code": "Low level programming.pdf",
    "operand": "Low level programming.pdf",
    "assembly language": "Low level programming.pdf",
    "index register": "Low level programming.pdf",
    "accumulator": "Low level programming.pdf",
    "label": "Low level programming.pdf",
    "acc": "Low level programming.pdf",
    "ldm": "Low level programming.pdf",
    "ldd": "Low level programming.pdf",
    "sto": "Low level programming.pdf",
    "cmp": "Low level programming.pdf",
    "jpe": "Low level programming.pdf",
    "jmp": "Low level programming.pdf",
    "ldi": "Low level programming.pdf",
    "add": "Low level programming.pdf",
    "addressing mode": "Low level programming.pdf",
    "content": "Low level programming.pdf",

    "clause": "Declarative programming.pdf",
    "knowledge base": "Declarative programming.pdf",
    "rule": "Declarative programming.pdf",
    "list": "Declarative programming.pdf",
    "fact": "Declarative programming.pdf",
    "predicate": "Declarative programming.pdf",
    "declarative": "Declarative programming.pdf",
    "goal": "Declarative programming.pdf",
    "additional clause": "Declarative programming.pdf",
    "if": "Declarative programming.pdf",

    "random file": "File processing and exception handling.pdf",
    "exception": "File processing and exception handling.pdf",
    "handling": "File processing and exception handling.pdf",

    "binary tree": "Algorithms.pdf",
    "linked list": "Algorithms.pdf",
    "left": "Algorithms.pdf",
    "right": "Algorithms.pdf",
    "pointer": "Algorithms.pdf",
    "null": "Algorithms.pdf",
    "node": "Algorithms.pdf",
    "rootptr": "Algorithms.pdf",
    "freeptr": "Algorithms.pdf",
    "1d": "Algorithms.pdf",
    "index": "Algorithms.pdf",
    "2d": "Algorithms.pdf",
    "linear search": "Algorithms.pdf",
    "bubble sort": "Algorithms.pdf",
    "insertion sort": "Algorithms.pdf",
    "binary search": "Algorithms.pdf",
    "abstract data type": "Algorithms.pdf",
    "stack": "Algorithms.pdf",
    "top": "Algorithms.pdf",
    "push": "Algorithms.pdf",
    "pop": "Algorithms.pdf",
    "queue": "Algorithms.pdf",
    "front": "Algorithms.pdf",
    "graph": "Algorithms.pdf",
    "hash": "Algorithms.pdf",
    "dictionary": "Algorithms.pdf",
    "dictionaries": "Algorithms.pdf",
    "big o notation": "Algorithms.pdf",
    "complexity": "Algorithms.pdf",
    "adt": "Algorithms.pdf",
    "user defined": "Algorithms.pdf",
    "endif": "Algorithms.pdf",
    "pseudocode": "Algorithms.pdf",

}
keyword_map4 = {
    "object": "OOP.pdf",
    "class": "OOP.pdf",
    "private": "OOP.pdf",
    "public": "OOP.pdf",
    "inheritance": "OOP.pdf",
    "aggregation": "OOP.pdf",
    "containment": "OOP.pdf",
    "OOP.pdf": "OOP.pdf",
    "attribute": "OOP.pdf",
    "method": "OOP.pdf",
    "encapsulation": "OOP.pdf",
    "getter": "OOP.pdf",
    "setter": "OOP.pdf",
    "constructor": "OOP.pdf",
    "instantiate": "OOP.pdf",
    "abstract": "OOP.pdf",
    "derived class": "OOP.pdf",
    "polymorphism": "OOP.pdf",

    "recursively": "Recursion & Arrays.pdf",
    "recursion": "Recursion & Arrays.pdf",
    "recursive": "Recursion & Arrays.pdf",
    "iterative": "Recursion & Arrays.pdf",
    "base case": "Recursion & Arrays.pdf",
    "general case": "Recursion & Arrays.pdf",
    "array": "Recursion & Arrays.pdf",
    "bubble sort": "Recursion & Arrays.pdf",
    "insertion sort": "Recursion & Arrays.pdf",

    "binary tree": "Binary Tree & Linked list.pdf",
    "linked list": "Binary Tree & Linked list.pdf",
    "left": "Binary Tree & Linked list.pdf",
    "right": "Binary Tree & Linked list.pdf",
    "node": "Binary Tree & Linked list.pdf",
    "rootptr": "Binary Tree & Linked list.pdf",
    "freeptr": "Binary Tree & Linked list.pdf",

    "stack": "Stack.pdf",
    "topofstack": "Stack.pdf",
    "push": "Stack.pdf",
    "pop": "Stack.pdf",

    "txt": "Text file.pdf",
    "text file": "Text file.pdf",
    "readdata()": "Text file.pdf",
    "insertdata()": "Text file.pdf",
    "getrecord": "Text file.pdf",

}


def get_chapter_list(paper):
    """
    Returns the list of chapter .docx filenames available for the given paper,
    derived from the keyword maps used during extraction. The .pdf -> .docx
    transform mirrors exactly what extract_questions() does when saving files,
    so the names returned here match the filenames that actually get written
    to the output folder.
    """
    if paper == "Paper 1":
        keyword_map = keyword_map1
    elif paper == "Paper 2":
        keyword_map = keyword_map2
    elif paper == "Paper 3":
        keyword_map = keyword_map3
    elif paper == "Paper 4":
        keyword_map = keyword_map4
    else:
        logging.warning(f"get_chapter_list: unknown paper '{paper}'")
        return []

    chapters = []
    for value in keyword_map.values():
        docx_name = value.replace(".pdf", ".docx")
        if docx_name not in chapters:
            chapters.append(docx_name)

    def chapter_sort_key(name):
        match = re.search(r"CH\s*#\s*c?(\d+)", name)
        return int(match.group(1)) if match else 9999

    chapters.sort(key=chapter_sort_key)
    return chapters


def _detect_footer_top(page, fallback_margin=80):
    """
    Finds the actual top y-coordinate of the page footer (the small
    '© UCLES ... [Turn over]' line) instead of assuming a fixed 80pt gap.
    A fixed margin cuts off legitimate content (e.g. tables) on any page
    where they run close to the bottom of the page.

    Only searches the bottom half of the page, because cover pages also
    have a "Cambridge International Examinations" logo/heading near the
    TOP of the page - without this restriction, that heading gets mistaken
    for the footer, producing a crop region that's inverted (bottom above
    top) and crashes pdfplumber. A safety clamp below guards against that
    happening from any other edge case too.
    """
    min_valid_top = barcode_y + 30  # never return something so small the crop region would be degenerate

    try:
        words = page.extract_words()
    except Exception:
        return page.height - fallback_margin

    bottom_half_start = page.height / 2
    footer_tops = [
        w["top"] for w in words
        if w["top"] >= bottom_half_start
        and ("UCLES" in w["text"].upper() or "CAMBRIDGE" in w["text"].upper())
    ]
    if footer_tops:
        candidate = max(0, min(footer_tops) - 8)  # small buffer above the footer text
        if candidate > min_valid_top:
            return candidate

    return page.height - fallback_margin


barcode_y = 55.07  # height to skip barcode

# Pages that contain no real question content — blank pages, and the
# "Question N begins on page M" / "Question N begins on the next page."
# notices CIE inserts between papers. These match any question number and
# either wording, so a page like "Question 9 begins on page 14." (not just
# "Question 1 begins on the next page.") is correctly skipped.
NOTICE_PAGE_PATTERNS = [
    r"BLANK PAGE",
    r"QUESTION\s+\d+\s+BEGINS\s+ON\s+(THE\s+NEXT\s+PAGE|PAGE\s+\d+)",
]


def add_question_to_docx(image_obj, chapter_docx_path, basename, add_question_num):
    # Open existing Word file or create new
    if os.path.exists(chapter_docx_path):
        doc = Document(chapter_docx_path)
    else:
        doc = Document()

    # Count existing questions by Heading 2 style
    question_no = sum(1 for p in doc.paragraphs if p.style.name == "Heading 2") + 1

    # Add header paragraph (question number)
    if add_question_num:
        heading = doc.add_paragraph(f"Q: {question_no} - ({basename})", style="Heading 2")
        heading.paragraph_format.space_after = Pt(13)

    # Save image to bytes
    img_bytes = io.BytesIO()
    image_obj.save(img_bytes, format="PNG")
    img_bytes.seek(0)

    # Scale image to fit page width (~6 inches)
    width_px, height_px = image_obj.size
    dpi = image_obj.info.get("dpi", (96, 96))[0] or 96
    width_in = width_px / dpi
    height_in = height_px / dpi

    target_width_in = 6.0
    scale_factor = target_width_in / width_in
    target_height_in = height_in * scale_factor
    doc.add_picture(img_bytes, width=Inches(target_width_in), height=Inches(target_height_in))

    # Add spacing after each question
    doc.add_paragraph("")

    # Save document
    doc.save(chapter_docx_path)
    print(f"[INFO] Added question {question_no} to {chapter_docx_path}")


def get_last_question_no(chapter_pdf_path):
    """Extract the last question number from an existing PDF, if any."""
    if not os.path.exists(chapter_pdf_path):
        return None

    doc = fitz.open(chapter_pdf_path)
    last_page = doc[-1]
    text = last_page.get_text("text")
    doc.close()

    # Look for "Answer X" pattern
    match = re.search(r"Answer\s+(\d+)", text)
    if match:
        return int(match.group(1))
    return None


def normalize_qnum(qnum):
    """Extract numeric part from labels like 'Q1', 'Q01', 'Q1a'."""
    match = re.search(r"\d+", str(qnum))
    return int(match.group(0)) if match else None


def get_last_answer_no(pdf_path):
    """Return the last answer number found in an existing PDF, or None."""
    if not os.path.exists(pdf_path):
        return None
    doc = fitz.open(pdf_path)
    last_page = doc[-1]
    text = last_page.get_text("text")
    doc.close()
    match = re.search(r"Q:\s*(\d+)", text)
    return int(match.group(1)) if match else None


def scale_image_to_page(image_obj, page_width, margin=40, top=100):
    """Scale image proportionally to fit page width with margins."""
    w, h = image_obj.size
    scale = (page_width - 2 * margin) / w
    new_w, new_h = int(w * scale), int(h * scale)
    return fitz.Rect(margin, top, margin + new_w, top + new_h)


def add_answer_to_chapter(image_obj, chapter_pdf_path, basename=""):
    """Add an answer image with header text to a chapter PDF, continuing numbering."""
    os.makedirs(os.path.dirname(chapter_pdf_path), exist_ok=True)

    # Open existing PDF or create new
    if os.path.exists(chapter_pdf_path):
        doc = fitz.open(chapter_pdf_path)
        new_file = False
        last_qnum = get_last_answer_no(chapter_pdf_path)
    else:
        doc = fitz.open()
        new_file = True
        last_qnum = None

    # Decide current question number
    if last_qnum is not None:
        current_qnum = last_qnum + 1
    else:
        current_qnum = 1

    # Convert PIL image to bytes
    img_bytes = io.BytesIO()
    image_obj.save(img_bytes, format="PNG")
    img_bytes = img_bytes.getvalue()

    # Create new A4 page
    rect = fitz.Rect(0, 0, 595, 842)
    page = doc.new_page(width=rect.width, height=rect.height)

    # Insert header text
    if add_answer_no == True:
        header_rect = fitz.Rect(0, 20, rect.width, 60)
        page.insert_textbox(
            header_rect,
            f"Q: {current_qnum} - ({basename})",
            fontsize=16,
            fontname="helvetica",
            align=1,
            color=(0, 0, 0)
        )

    # Insert scaled image
    img_rect = scale_image_to_page(image_obj, page_width=rect.width)
    page.insert_image(img_rect, stream=img_bytes)

    # Footer with page number
    footer_rect = fitz.Rect(0, 800, rect.width, 842)
    page.insert_textbox(
        footer_rect,
        f"Page {doc.page_count}",
        fontsize=10,
        fontname="helvetica",
        align=2,
        color=(0, 0, 0)
    )

    # Save and close
    try:
        doc.save(chapter_pdf_path, incremental=not new_file, encryption=fitz.PDF_ENCRYPT_KEEP)
    finally:
        doc.close()

    print(f"[INFO] Added answer {current_qnum} to {chapter_pdf_path}")


def is_question_header(text):
    stripped = text.lstrip()
    space_depth = len(text) - len(stripped)

    # Only consider lines where the number starts within the first few spaces
    if space_depth > 3:
        return False

    text = text.strip()

    # Reject short lines
    if len(text.split()) < 4:
        return False

    if re.match(r"^0\b", text):  # line starts with zero
        return False

    # Reject patterns like "5 1 2 3 (Hardware) 8"
    if re.match(r"^\d+\s+\d+", text):
        return False

    # Reject lines like "1. something" — structured or procedural
    if re.match(r"^\d{1,2}\.", text):
        return False

    # Reject table-style assignments, like "1 Score = 65"
    if re.match(r"^\d{1,2}\s+\w+\s*=\s*\w+", text):
        return False

    # Reject dotted filler lines like "2 ..........."
    if re.match(r"^\d{1,2}[\.\)]?\s*[\.\s]*$", text):
        return False

    # Accept composite headers like "2 (a) Describe..."
    if re.match(r"^\d{1,2}(\s*\([a-zA-Z]\))+\s+\w+", text):
        return True

    # Accept basic headers like "3 (Hardware) Describe the concept..."
    return bool(re.match(r"^\d{1,2}[\)]?\s+\w+", text))


def is_sub_question(text):
    text = text.strip()
    first_token = text.split()[0] if text else ""
    return bool(re.fullmatch(r"\(\s*[ivxlcdm]+\s*\)", first_token, re.IGNORECASE))


def is_alphabet_subpart(text):
    text = text.strip()
    first_token = text.split()[0] if text else ""
    return bool(re.fullmatch(r"\([a-zA-Z]\)", first_token))


def keyword_in_text(keyword, text):
    # normalize spaces and lowercase
    text = re.sub(r"\s+", " ", text.lower())
    keyword = keyword.lower()

    # allow small spacing variations between words
    if " " in keyword:
        words = keyword.split()
        pattern = r"\b" + words[0] + r"\s*" + words[1] + r"\b"
        return re.search(pattern, text) is not None
    else:
        pattern = r"\b" + re.escape(keyword) + r"\b"
        return re.search(pattern, text) is not None


"""def keyword_in_text(keyword, text):
    pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
    return re.search(pattern, text.lower()) is not None"""


def assign_best_chapter(question_text, keyword_map, chapter_folder, threshold=2, margin=1):
    """
    Score each chapter by keyword hits in question_text and pick the single best one.

    threshold: minimum score the top chapter must reach to be considered a valid match.
    margin:    minimum score difference the top chapter must have over the
               second-best chapter to be considered unambiguous. If the top two
               chapters are within `margin` of each other, we refuse to guess
               and return None instead of silently picking one.
    """
    global paper_label, best_chapter

    chapter_scores = defaultdict(int)
    chapter_hits = defaultdict(list)  # which keywords matched, for debugging/logging

    for keyword, chapter_file in keyword_map.items():
        if keyword_in_text(keyword, question_text):
            chapter_scores[chapter_file] += 1
            chapter_hits[chapter_file].append(keyword)

    if not chapter_scores:
        return None

    # Sort chapters by score, descending
    ranked = sorted(chapter_scores.items(), key=lambda kv: kv[1], reverse=True)

    top_chapter, top_score = ranked[0]

    # Must clear the minimum threshold at all
    if top_score < threshold:
        return None

    # If there's a second-place chapter, check the margin between them
    if len(ranked) > 1:
        second_chapter, second_score = ranked[1]
        if (top_score - second_score) < margin:
            print(
                f"[WARN] Ambiguous chapter assignment "
                f"({top_chapter}={top_score} vs {second_chapter}={second_score}); "
                f"matched keywords: {top_chapter}->{chapter_hits[top_chapter]}, "
                f"{second_chapter}->{chapter_hits[second_chapter]}. Skipping assignment."
            )
            return None

    best_chapter = top_chapter
    return os.path.join(chapter_folder, best_chapter)


def extract_questions(filenames, mode, chapter_folder, ui=None):
    global add_answer_no, add_question_num, PreviousQuestionNo, last_chapter_path, paper_label, basename, best_chapter

    if isinstance(filenames, str):
        filenames = [filenames]

    results = {}

    # --- Handle mode logic ---
    if mode == "new":
        for f in os.listdir(chapter_folder):
            path = os.path.join(chapter_folder, f)
            if os.path.isfile(path):
                os.remove(path)
            elif os.path.isdir(path):
                import shutil
                shutil.rmtree(path)
        print(f"[INFO] Starting NEW extraction in {chapter_folder}")
    else:
        print(f"[INFO] Using EXISTING extraction in {chapter_folder}")

    for filename in filenames:
        add_question_num = False
        PreviousQuestionNo = None
        last_chapter_path = None
        questions = []

        basename = os.path.splitext(os.path.basename(filename))[0]

        paper_code = basename.split("_")[-1]
        paper_category = basename.split("_")[0]
        if paper_category in ("9618", "9608"):

            if paper_code in ("11", "12", "13"):
                paper_label = "Paper 1"
            elif paper_code in ("21", "22", "23"):
                paper_label = "Paper 2"
            elif paper_code in ("31", "32", "33"):
                paper_label = "Paper 3"
            elif paper_code in ("41", "42", "43"):
                paper_label = "paper 4"
            else:
                print(f"[WARN] Unrecognized paper code '{paper_code}' for {basename} — no paper_label set")
                paper_label = None
        else:
            raise InvalidPaperError(
                f"'{basename}' is not a recognized A Level paper (expected filename "
                f"starting with 9618 or 9608)."
            )
        os.makedirs(chapter_folder, exist_ok=True)

        if paper_label == "Paper 1":
            keyword_map = keyword_map1
        elif paper_label == "Paper 2":
            keyword_map = keyword_map2
        elif paper_label == "Paper 3":
            keyword_map = keyword_map3
        elif paper_label == "paper 4":
            keyword_map = keyword_map4
        else:
            keyword_map = {}

        # --- Buffer for the question currently being accumulated ---
        # We hold every block (image + OCR text) belonging to the same question
        # number until the whole question has been read (across sub-parts, and
        # even across a page boundary), THEN decide the chapter once from the
        # combined text, then write every block together under that decision.
        pending_blocks = []  # list of (cropped_img, add_heading_flag)
        pending_text_parts = []  # OCR'd text fragments for the current question
        pending_qnum = None
        heading_pending = False  # True until the first block of a new question is stored

        def flush_pending():
            nonlocal pending_blocks, pending_text_parts, pending_qnum, heading_pending
            if not pending_blocks:
                return
            combined_text = " ".join(pending_text_parts)
            chapter_paths = assign_all_chapters(combined_text, keyword_map, chapter_folder, threshold=2)
            safe_qnum = f"Q{pending_qnum}"
            if chapter_paths:
                for cropped_img, add_heading in pending_blocks:
                    for chapter_path in chapter_paths:
                        docx_path = chapter_path.replace(".pdf", ".docx")
                        add_question_to_docx(cropped_img, docx_path, basename, add_heading)
                        if add_heading:
                            global add_answer_no
                            add_answer_no = True
                        print(f"Added {safe_qnum} to {os.path.basename(docx_path)}")
                for chapter_path in chapter_paths:
                    questions.append(os.path.basename(chapter_path).replace(".pdf", ".docx"))
            else:
                print(f"No chapter assigned for {safe_qnum}, skipping...")
            pending_blocks = []
            pending_text_parts = []
            pending_qnum = None
            heading_pending = False

        with pdfplumber.open(filename) as pdf:
            for page_num, page in enumerate(pdf.pages[1:], start=2):
                add_question_num = False
                noQuestionNo = False

                footer_top = _detect_footer_top(page)
                footer_top = max(footer_top, barcode_y + 30)  # guarantee a valid (non-inverted) crop region no matter what
                cropped_page = page.crop((25, barcode_y + 10, page.width, footer_top))
                lines = cropped_page.extract_text(layout=True)
                if not lines:
                    continue

                full_text = page.extract_text()
                if full_text:
                    normalized_text = re.sub(r"\s+", " ", full_text.upper())
                    if any(re.search(pat, normalized_text) for pat in NOTICE_PAGE_PATTERNS):
                        continue

                text_lines = lines.split("\n")
                line_height = 16
                table_boxes = [tbl.bbox for tbl in page.find_tables()]
                table_regions = [(top, bottom) for (_, top, _, bottom) in table_boxes]

                headers = []
                header_flag = False
                for line_num, line in enumerate(text_lines):
                    line_y = line_num * line_height
                    if any(top <= line_y <= bottom for top, bottom in table_regions):
                        continue
                    if is_question_header(line):
                        headers.append((line_num, line))
                        header_flag = True
                    elif line_num <= 3 and is_alphabet_subpart(line) and not header_flag:
                        headers.append((line_num, line))
                        break
                    elif line_num <= 3 and is_sub_question(line) and not header_flag:
                        headers.append((line_num, line))
                        break

                headers.append((len(text_lines), "END"))

                if len(headers) == 1 and headers[0][1] == "END" and PreviousQuestionNo:
                    headers.insert(0, (0, f"{PreviousQuestionNo} (continued)"))
                    noQuestionNo = True

                image = cropped_page.to_image(resolution=300)
                line_height = image.original.height / max(len(text_lines), 1)

                qnum_match = re.match(r"^\s*(\d+)", headers[0][1])

                if qnum_match and not noQuestionNo:
                    question_num = qnum_match.group(1)
                    if question_num != PreviousQuestionNo:
                        # Brand-new question — flush the previous one (this
                        # decides its chapter from ALL its accumulated blocks)
                        # before we start collecting for the new one.
                        flush_pending()
                        pending_qnum = question_num
                        heading_pending = True
                    PreviousQuestionNo = question_num
                    subpartFlag = False
                else:
                    question_num = PreviousQuestionNo if PreviousQuestionNo else str(0 + 1)
                    subpartFlag = bool(PreviousQuestionNo)
                    if pending_qnum is None:
                        pending_qnum = question_num
                        heading_pending = True

                safe_qnum = f"Q{question_num}"

                for i in range(len(headers) - 1):
                    line_start, line_end = headers[i][0], headers[i + 1][0]
                    buffer = 40
                    top_y = max(0, int(line_start * line_height) - buffer)
                    bottom_y = int(line_end * line_height)
                    if bottom_y <= top_y:
                        continue
                    top_crop = max(0, top_y - 15)

                    cropped_img = image.original.crop(
                        (0, top_crop, image.original.width - 80, bottom_y)
                    )

                    # OCR this block and add its text to the current question's buffer.
                    # No chapter decision happens here anymore — only at flush_pending().
                    # Tesseract path is already configured at import time via _configure_tesseract()
                    block_text = pytesseract.image_to_string(cropped_img).lower()
                    block_text = re.sub(r"[^\w\s]", "", block_text)
                    pending_text_parts.append(block_text)

                    # Exactly one heading per question, on the very first block
                    # we ever collect for it — regardless of how many sub-parts
                    # or pages it spans.
                    add_heading = heading_pending
                    if heading_pending:
                        heading_pending = False
                    pending_blocks.append((cropped_img, add_heading))

        # Flush whatever's left buffered at the end of this file
        flush_pending()

        results[os.path.basename(filename)] = questions

    return results


from docx import Document
from docxcompose.composer import Composer


def assign_all_chapters(question_text, keyword_map, chapter_folder, threshold=2, margin_ratio=0.5):
    chapter_scores = defaultdict(int)
    for keyword, chapter_file in keyword_map.items():
        if keyword_in_text(keyword, question_text):
            chapter_scores[chapter_file] += 1

    if not chapter_scores:
        return []

    top_score = max(chapter_scores.values())
    selected_chapters = [
        os.path.join(chapter_folder, chapter_file)
        for chapter_file, score in chapter_scores.items()
        if score >= threshold and score >= top_score * margin_ratio
    ]
    return selected_chapters


def generate_test(chapter_paths, output_path):
    test_title = "A Level Test"
    # Create a new master document
    master = Document()
    master.add_heading(test_title, level=0)

    composer = Composer(master)

    for chapter_path in chapter_paths:
        try:
            chapter_doc = Document(chapter_path)
            composer.append(chapter_doc)
            # Optional: add a page break between chapters
            master.add_page_break()
            print(f"[INFO] Added {chapter_path} to test")
        except Exception as e:
            print(f"[ERROR] Could not add {chapter_path}: {e}")

    composer.save(output_path)
    print(f"[INFO] Test created at {output_path}")


from docx import Document


def renumber_questions(docx_path):
    doc = Document(docx_path)
    qnum = 1
    for para in doc.paragraphs:
        if para.style.name == "Heading 2" and para.text.strip().startswith("Q:"):
            # Replace with just "Q: <number>"
            new_text = f"Q: {qnum}"

            # Update only the first run to preserve formatting
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.text = new_text

            qnum += 1

    doc.save(docx_path)
    print(f"[INFO] Renumbered questions (removed paper refs) in {docx_path}")

